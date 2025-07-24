import os
import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain.docstore.document import Document
from langchain.llms import HuggingFacePipeline
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
import tempfile
from docx import Document as DocxDocument
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline
import glob
import re
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import fitz  # PyMuPDF for better PDF processing
import pandas as pd
import os

static_dir = os.path.join(os.path.dirname(__file__), "static")
os.makedirs(static_dir, exist_ok=True)

class RAGSystem:
    def __init__(self):
        self.embeddings = None
        self.vectorstore = None
        self.qa_chain = None
        self.setup_embeddings()
        self.setup_llm()
        self.preload_sample_documents()
    
    def setup_embeddings(self):
        """Initialize embedding model"""
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'}
        )
    
    def setup_llm(self):
        """Initialize open-source LLM"""
        model_name = "microsoft/DialoGPT-medium"
        
        # Use a lightweight model for demonstration
        tokenizer = AutoTokenizer.from_pretrained(model_name)
        model = AutoModelForCausalLM.from_pretrained(model_name)
        
        # Create pipeline
        pipe = pipeline(
            "text-generation",
            model=model,
            tokenizer=tokenizer,
            max_length=512,
            temperature=0.7,
            do_sample=True,
            device=-1  # CPU
        )
        
        self.llm = HuggingFacePipeline(pipeline=pipe)
    
    def load_documents(self, uploaded_files):
        """Load and process uploaded documents"""
        documents = []
        
        for uploaded_file in uploaded_files:
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_file_path = tmp_file.name
            
            try:
                if uploaded_file.name.endswith('.pdf'):
                    loader = PyPDFLoader(tmp_file_path)
                    docs = loader.load()
                elif uploaded_file.name.endswith('.txt'):
                    loader = TextLoader(tmp_file_path)
                    docs = loader.load()
                elif uploaded_file.name.endswith('.docx'):
                    doc = DocxDocument(tmp_file_path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    docs = [Document(page_content=text, metadata={"source": uploaded_file.name})]
                else:
                    st.warning(f"Unsupported file type: {uploaded_file.name}")
                    continue
                
                documents.extend(docs)
            finally:
                os.unlink(tmp_file_path)
        
        return documents
    
    def preload_sample_documents(self):
        """Preload documents from Sample datasets folder"""
        sample_folder = os.path.join(os.path.dirname(__file__), "Sample datasets")
        
        if not os.path.exists(sample_folder):
            st.warning(f"Sample datasets folder not found at: {sample_folder}")
            return
        
        # Find all PDF files in the sample folder
        pdf_files = glob.glob(os.path.join(sample_folder, "*.pdf"))
        
        if not pdf_files:
            st.info("No PDF files found in Sample datasets folder")
            return
        
        try:
            documents = []
            for pdf_file in pdf_files:
                loader = PyPDFLoader(pdf_file)
                docs = loader.load()
                # Add filename to metadata
                for doc in docs:
                    doc.metadata["source"] = os.path.basename(pdf_file)
                documents.extend(docs)
            
            if documents:
                # Create vector store   
                vectorstore = self.create_vectorstore(documents)
                if vectorstore:
                    # Setup QA chain
                    qa_chain = self.setup_qa_chain()
                    if qa_chain:
                        st.session_state.documents_processed = True
                        st.session_state.preloaded_docs = len(documents)
                        st.success(f"✅ Preloaded {len(documents)} documents from Sample datasets!")
        except Exception as e:
            st.error(f"Error preloading sample documents: {str(e)}")
    
    def smart_chunking(self, documents):
        """Enhanced chunking strategy for different document layouts"""
        chunks = []
        
        for doc in documents:
            content = doc.page_content
            metadata = doc.metadata
            
            # Detect and handle tables
            table_pattern = r'\|.*\|.*\n'
            tables = re.findall(table_pattern, content, re.MULTILINE)
            
            # Detect structured content (headers, lists)
            header_pattern = r'^#{1,6}\s+.*$|^[A-Z][^\n]*:$'
            list_pattern = r'^\s*[-*•]\s+.*$|^\s*\d+\.\s+.*$'
            
            # Split by semantic boundaries
            sections = re.split(r'\n\s*\n', content)
            
            for i, section in enumerate(sections):
                if len(section.strip()) < 50:  # Skip very short sections
                    continue
                
                # Determine chunk type
                chunk_type = "text"
                if any(table in section for table in tables):
                    chunk_type = "table"
                elif re.search(header_pattern, section, re.MULTILINE):
                    chunk_type = "header"
                elif re.search(list_pattern, section, re.MULTILINE):
                    chunk_type = "list"
                
                # Create enhanced metadata
                enhanced_metadata = metadata.copy()
                enhanced_metadata.update({
                    "chunk_type": chunk_type,
                    "chunk_index": i,
                    "section_context": self.extract_context(sections, i)
                })
                
                # Create overlapping chunks for better context
                if len(section) > 800:
                    sub_chunks = self.create_overlapping_chunks(section, 600, 100)
                    for j, sub_chunk in enumerate(sub_chunks):
                        sub_metadata = enhanced_metadata.copy()
                        sub_metadata["sub_chunk_index"] = j
                        chunks.append(Document(page_content=sub_chunk, metadata=sub_metadata))
                else:
                    chunks.append(Document(page_content=section, metadata=enhanced_metadata))
        
        return chunks
    
    def extract_context(self, sections, current_index):
        """Extract surrounding context for better understanding"""
        context = []
        # Add previous section title/header if available
        if current_index > 0:
            prev_section = sections[current_index - 1][:100]
            context.append(f"Previous: {prev_section}")
        
        # Add next section title/header if available
        if current_index < len(sections) - 1:
            next_section = sections[current_index + 1][:100]
            context.append(f"Next: {next_section}")
        
        return " | ".join(context)
    
    def create_overlapping_chunks(self, text, chunk_size, overlap):
        """Create overlapping chunks with sentence boundary awareness"""
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk + sentence) > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                # Keep overlap
                overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                current_chunk = overlap_text + " " + sentence
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def create_vectorstore(self, documents):
        """Create vector store with enhanced chunking"""
        if not documents:
            return None
        
        # Use smart chunking instead of basic text splitting
        texts = self.smart_chunking(documents)
        
        # Create vector store
        self.vectorstore = FAISS.from_documents(texts, self.embeddings)
        
        return self.vectorstore
    
    def contextual_retriever(self, query, k=5):
        """Enhanced retrieval with contextual understanding"""
        if not self.vectorstore:
            return []
        
        # Get initial candidates
        initial_docs = self.vectorstore.similarity_search_with_score(query, k=k*2)
        
        # Re-rank based on context and relevance
        reranked_docs = []
        query_embedding = self.embeddings.embed_query(query)
        
        for doc, score in initial_docs:
            # Calculate contextual relevance
            context_score = self.calculate_context_relevance(doc, query)
            
            # Combine similarity and context scores
            combined_score = (1 - score) * 0.7 + context_score * 0.3
            
            reranked_docs.append((doc, combined_score))
        
        # Sort by combined score and return top k
        reranked_docs.sort(key=lambda x: x[1], reverse=True)
        return [doc for doc, _ in reranked_docs[:k]]
    
    def calculate_context_relevance(self, doc, query):
        """Calculate contextual relevance score"""
        content = doc.page_content.lower()
        query_lower = query.lower()
        
        # Keyword overlap score
        query_words = set(re.findall(r'\w+', query_lower))
        content_words = set(re.findall(r'\w+', content))
        keyword_overlap = len(query_words.intersection(content_words)) / len(query_words) if query_words else 0
        
        # Chunk type relevance
        chunk_type = doc.metadata.get('chunk_type', 'text')
        type_bonus = 0.1 if chunk_type in ['table', 'header'] and any(word in query_lower for word in ['table', 'data', 'summary', 'overview']) else 0
        
        # Section context bonus
        section_context = doc.metadata.get('section_context', '')
        context_bonus = 0.05 if any(word in section_context.lower() for word in query_words) else 0
        
        return keyword_overlap + type_bonus + context_bonus
    
    def setup_qa_chain(self):
        """Setup QA chain with enhanced contextual understanding"""
        if not self.vectorstore:
            return None

        prompt_template = """You are an intelligent assistant that answers questions based on provided context. 
        Analyze the context carefully, considering different types of content (text, tables, headers, lists).

        Context Information:
        {context}

        Question: {question}

        Instructions:
        1. Focus only on information directly relevant to the question
        2. If the context contains tables or structured data, interpret them appropriately
        3. Combine information from multiple sources when necessary
        4. If you cannot find relevant information, clearly state that
        5. Provide specific details and cite sources when possible

        Answer: """

        PROMPT = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )

        # Use the built-in retriever from vectorstore
        retriever = self.vectorstore.as_retriever(search_kwargs={"k": 5})

        self.qa_chain = RetrievalQA.from_chain_type(
            llm=self.llm,
            chain_type="stuff",
            retriever=retriever,
            chain_type_kwargs={"prompt": PROMPT},
            return_source_documents=True
        )

        return self.qa_chain
    
    def query(self, question):
        """Enhanced query processing with context analysis"""
        if not self.qa_chain:
            return "Please upload documents first to create the knowledge base."
        
        try:
            # Analyze query for better retrieval
            query_analysis = self.analyze_query(question)
            
            # Get contextually relevant documents
            relevant_docs = self.contextual_retriever(question, k=3)
            
            # Format context with metadata
            formatted_context = self.format_context_with_metadata(relevant_docs)
            
            # Generate response
            result = self.qa_chain({"query": question})
            
            # Enhance result with relevance scores and source info
            if isinstance(result, dict) and 'source_documents' in result:
                result['relevance_analysis'] = query_analysis
                result['context_types'] = [doc.metadata.get('chunk_type', 'text') for doc in relevant_docs]
            
            return result
        except Exception as e:
            return f"Error processing query: {str(e)}"
    
    def analyze_query(self, query):
        """Analyze query to understand intent and what type of context needs"""
        query_lower = query.lower()
        
        analysis = {
            'intent': 'general',
            'needs_tables': any(word in query_lower for word in ['table', 'data', 'numbers', 'statistics', 'compare']),
            'needs_summary': any(word in query_lower for word in ['summary', 'overview', 'main points', 'key']),
            'needs_specific': any(word in query_lower for word in ['specific', 'detail', 'exactly', 'precisely']),
            'question_type': 'what' if 'what' in query_lower else 'how' if 'how' in query_lower else 'why' if 'why' in query_lower else 'other'
        }
        
        return analysis
    
    def format_context_with_metadata(self, docs):
        """Format retrieved documents with their metadata for better context"""
        formatted_context = []
        
        for doc in docs:
            chunk_type = doc.metadata.get('chunk_type', 'text')
            source = doc.metadata.get('source', 'unknown')
            
            context_header = f"[{chunk_type.upper()} from {source}]"
            formatted_context.append(f"{context_header}\n{doc.page_content}")
        
        return "\n\n".join(formatted_context)

def main():
    st.set_page_config(page_title="RAG System", page_icon="🤖", layout="wide")
    
    st.title("🤖 RAG (Retrieval-Augmented Generation) System")
    st.markdown("Upload documents and ask questions based on their content!")
    
    # Initialize RAG system
    if 'rag_system' not in st.session_state:
        with st.spinner("Initializing RAG system..."):
            st.session_state.rag_system = RAGSystem()
    
    # Sidebar for document upload
    with st.sidebar:
        st.header("📁 Document Management")
        
        # Show preloaded documents status
        if hasattr(st.session_state, 'preloaded_docs'):
            st.info(f"📚 {st.session_state.preloaded_docs} sample documents loaded")
        
        st.subheader("Upload Additional Documents")
        uploaded_files = st.file_uploader(
            "Choose files",
            accept_multiple_files=True,
            type=['pdf', 'txt', 'docx']
        )
        
        if uploaded_files:
            if st.button("Process Additional Documents"):
                with st.spinner("Processing documents..."):
                    # Load documents
                    documents = st.session_state.rag_system.load_documents(uploaded_files)
                    
                    if documents:
                        # Create vector store (this will replace existing one)
                        vectorstore = st.session_state.rag_system.create_vectorstore(documents)
                        
                        if vectorstore:
                            # Setup QA chain
                            qa_chain = st.session_state.rag_system.setup_qa_chain()
                            
                            if qa_chain:
                                st.success(f"Successfully processed {len(documents)} additional documents!")
                                st.session_state.documents_processed = True
                            else:
                                st.error("Failed to setup QA chain")
                        else:
                            st.error("Failed to create vector store")
                    else:
                        st.error("No documents were successfully loaded")
    
    # Main chat interface
    st.header("💬 Chat Interface")
    
    # Initialize chat history
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            
            # Display source information if available
            if "sources" in message and message["sources"]:
                with st.expander("📚 Sources & Context Analysis"):
                    for i, source in enumerate(message["sources"]):
                        chunk_type = source.metadata.get('chunk_type', 'text')
                        source_file = source.metadata.get('source', 'unknown')
                        
                        st.write(f"**Source {i+1}** ({chunk_type.title()} from {source_file})")
                        st.write(source.page_content[:300] + "..." if len(source.page_content) > 300 else source.page_content)
                        st.write("---")
    
    # Chat input
    if prompt := st.chat_input("Ask a question about your documents..."):
        # Add user message
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Generate response
        with st.chat_message("assistant"):
            with st.spinner("Analyzing documents and generating response..."):
                response = st.session_state.rag_system.query(prompt)
                
                if isinstance(response, dict):
                    # Display main answer
                    answer = response.get('result', 'No answer generated')
                    st.markdown(answer)
                    
                    # Display query analysis
                    if 'relevance_analysis' in response:
                        analysis = response['relevance_analysis']
                        with st.expander("🔍 Query Analysis"):
                            col1, col2 = st.columns(2)
                            with col1:
                                st.write(f"**Question Type:** {analysis['question_type']}")
                                st.write(f"**Intent:** {analysis['intent']}")
                            with col2:
                                st.write(f"**Needs Tables:** {'Yes' if analysis['needs_tables'] else 'No'}")
                                st.write(f"**Needs Summary:** {'Yes' if analysis['needs_summary'] else 'No'}")
                    
                    # Display context types found
                    if 'context_types' in response:
                        context_types = response['context_types']
                        st.info(f"📊 Retrieved content types: {', '.join(set(context_types))}")
                    
                    # Store sources for display
                    sources = response.get('source_documents', [])
                    
                    # Add assistant message with sources
                    st.session_state.messages.append({
                        "role": "assistant", 
                        "content": answer,
                        "sources": sources
                    })
                else:
                    # Handle string response
                    st.markdown(str(response))
                    st.session_state.messages.append({"role": "assistant", "content": str(response)})
    
    # System information
    with st.expander("🔧 System Information"):
        st.write("**Enhanced Features:**")
        st.write("✅ Smart chunking with layout awareness")
        st.write("✅ Table and structured content detection")
        st.write("✅ Contextual understanding and re-ranking")
        st.write("✅ Multi-source information synthesis")
        st.write("✅ Query intent analysis")
        st.write("✅ Relevance-based filtering")
    
    # Instructions
    if not hasattr(st.session_state, 'documents_processed'):
        st.info("🤖 Sample documents are being loaded automatically. You can also upload additional documents in the sidebar!")
    
    # # Sample questions for preloaded documents
    # if hasattr(st.session_state, 'documents_processed') and st.session_state.documents_processed:
    #     st.header("💡 Sample Questions")
        
    #     col1, col2, col3 = st.columns(3)
        
    #     with col1:
    #         if st.button("📋 What are the main topics?"):
    #             st.session_state.messages.append({"role": "user", "content": "What are the main topics covered in the documents?"})
    #             st.rerun()
        
    #     with col2:
    #         if st.button("📊 Show me data/tables"):
    #             st.session_state.messages.append({"role": "user", "content": "Can you show me any tables or data from the documents?"})
    #             st.rerun()
        
    #     with col3:
    #         if st.button("📝 Summarize key points"):
    #             st.session_state.messages.append({"role": "user", "content": "Please summarize the key points from all documents."})
    #             st.rerun()
        
        # # Additional sample questions
        # st.subheader("More Questions")
        # sample_questions = [
        #     "What are the key features mentioned in the documents?",
        #     "Tell me about the main topics covered",
        #     "What is the RAG system architecture?",
        #     "Summarize the main points from the documents"
        # ]
        
        # cols = st.columns(2)
        # for i, question in enumerate(sample_questions):
        #     with cols[i % 2]:
        #         if st.button(question, key=f"sample_{i}"):
        #             # Add the question to chat
        #             st.session_state.messages.append({"role": "user", "content": question})
        #             st.rerun()

if __name__ == "__main__":
    main()